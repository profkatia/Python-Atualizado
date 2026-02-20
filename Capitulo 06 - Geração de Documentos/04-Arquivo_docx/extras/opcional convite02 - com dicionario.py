from docx import Document

def criar_convite(nome, idade):
    doc = Document()
    
    doc.add_heading('Convite de Aniversário', level=1)
    doc.add_paragraph("Você está convidado(a) para a festa de aniversário de:")
    doc.add_paragraph(f"Nome: {nome}")
    doc.add_paragraph(f"Idade: {idade} anos")

    nome_arquivo = f'convite_{nome}.docx'
    doc.save(nome_arquivo)

    print(f'Convite criado: {nome_arquivo}')


# 📌 Lista de convidados
convidados = [
    ("Ana", 25),
    ("Carlos", 30),
    ("Marina", 28),
    ("João", 35)
]

# Loop para gerar vários convites
for nome, idade in convidados:
    criar_convite(nome, idade)