import pandas as pd
from docx import Document

def criar_convite(nome, idade):
    doc = Document()
    
    doc.add_heading('Convite de Aniversário', level=1)
    doc.add_paragraph("Você está convidado(a) para a festa de aniversário de:")
    doc.add_paragraph(f"Nome: {nome}")
    doc.add_paragraph(f"Idade: {idade} anos")

    nome_arquivo = f'convite_{nome}.docx'
    doc.save(nome_arquivo)

    print(f'Convite criado: {nome_arquivo}')


# 📌 Lendo planilha
df = pd.read_excel("convidados.xlsx")

# Gerando convites automaticamente
for index, linha in df.iterrows():
    criar_convite(linha["Nome"], linha["Idade"])