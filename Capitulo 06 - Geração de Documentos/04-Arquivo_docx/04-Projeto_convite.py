from docx import Document
from docx.shared import Inches

def criar_convite(nome, idade):
    doc = Document()
    
    doc.add_heading('Convite de Aniversário', level=1)

    doc.add_paragraph("Você está convidado(a) para a festa de aniversário de:")
    doc.add_paragraph(f"Nome: {nome}")
    doc.add_paragraph(f"Idade: {idade} anos")

    nome_arquivo_docx = f'convite_{nome}.docx'

    doc.save(nome_arquivo_docx)

    print(f'Convite em DOCX para {nome} criado: {nome_arquivo_docx}')


# Programa principal
nome_convidado = input("Nome do convidado: ")
idade_convidado = int(input("Idade do convidado: "))

criar_convite(nome_convidado, idade_convidado)