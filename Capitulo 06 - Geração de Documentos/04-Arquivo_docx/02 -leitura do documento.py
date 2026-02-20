#fazer leitura do documento 

from docx import Document

doc = Document()
doc = Document('documento.docx')
for paragraph in doc.paragraphs:
    print(paragraph.text)     