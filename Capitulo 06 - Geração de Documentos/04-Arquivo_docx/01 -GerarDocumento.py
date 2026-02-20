#documetos docx
#pip install python-docx

from docx import Document

doc = Document() 

doc.add_heading('Titulo', level=1)
doc.add_paragraph('Este é um parágrafo')

doc.add_heading('Subtitulo', level=2)
doc.add_paragraph('Outro paragrafo')

doc.save('documento.docx')



  

