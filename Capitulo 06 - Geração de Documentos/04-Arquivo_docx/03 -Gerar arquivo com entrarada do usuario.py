#Gerar arquivo com entrarada do usuario
from docx import Document

doc = Document()    

titulo = input('Digite o titulo do documento: ')
doc.add_heading(titulo, level=1)    
paragrafo = input('Digite o parágrafo do documento: ')
doc.add_paragraph(paragrafo)
doc.save('documento_usuario.docx')